import tkinter as tk
from tkinter import scrolledtext
from transformers import AutoModelForCausalLM, AutoTokenizer
import torch
import re
import os
from datasets import load_dataset
from fpdf import FPDF
from docx import Document

# 檢查 CUDA 是否可用
DEVICE = "cuda" if torch.cuda.is_available() else "cpu"

# 模型與 Tokenizer 設定
MODEL_NAME = "microsoft/DialoGPT-medium"

# 清理文本函數
def clean_text(text):
    text = re.sub(r"<[^>]+>", "", text)  # 移除HTML標籤
    text = re.sub(r"[^\u4e00-\u9fa5a-zA-Z0-9，。！？：；「」『』]", "", text)  # 保留中文、英文、數字
    text = re.sub(r"\s+", " ", text)  # 移除多餘空白
    return text.strip()

# 生成 PDF
def save_as_pdf(cleaned_data):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=1, margin=15)
    pdf.set_font("Arial", size=12)

    for line in cleaned_data:
        pdf.multi_cell(0, 10, line, align="L")

    pdf.output("cleaned_data/cleaned_data.pdf")

# 生成 Word
def save_as_word(cleaned_data):
    doc = Document()
    doc.add_heading("清理後的資料", level=1)

    for line in cleaned_data:
        doc.add_paragraph(line)

    doc.save("cleaned_data/cleaned_data.docx")

# 生成 TXT
def save_as_txt(cleaned_data):
    with open("cleaned_data/cleaned_data.txt", "w", encoding="utf-8") as f:
        f.write("\n".join(cleaned_data))

# 載入與清理資料
def prepare_dataset():
    try:
        dataset = load_dataset("text", data_files={"train": "train_data.txt"})
        cleaned_data = [clean_text(line["text"]) for line in dataset["train"] if len(clean_text(line["text"])) > 10]

        # 建立輸出資料夾
        if not os.path.exists("cleaned_data"):
            os.makedirs("cleaned_data")

        # 保存為多種格式
        save_as_txt(cleaned_data)
        save_as_pdf(cleaned_data)
        save_as_word(cleaned_data)

        return True
    except Exception as e:
        response_text.insert(tk.END, f"資料處理失敗：{str(e)}\n")
        return False

# 載入模型
def load_model():
    global model, tokenizer
    tokenizer = AutoTokenizer.from_pretrained(MODEL_NAME)
    model = AutoModelForCausalLM.from_pretrained(MODEL_NAME, device_map="auto").to(DEVICE)

# 生成回應
def generate_response():
    user_input = user_input_text.get("1.0", tk.END).strip()
    if not user_input:
        return

    inputs = tokenizer(user_input, return_tensors="pt").to(DEVICE)
    output = model.generate(
        **inputs, max_new_tokens=200, do_sample=True, temperature=0.7
    )

    response = tokenizer.decode(output[0], skip_special_tokens=True)
    response_text.insert(tk.END, f"你: {user_input}\nLLM: {response}\n{'-'*40}\n")
    user_input_text.delete("1.0", tk.END)

# GUI 介面
root = tk.Tk()
root.title("簡單 LLM GUI")
root.geometry("600x600")

# 歡迎標題
welcome_label = tk.Label(root, text="簡單 LLM 聊天機器人", font=("Arial", 16))
welcome_label.pack(pady=10)

# 對話輸出框
response_text = scrolledtext.ScrolledText(root, wrap=tk.WORD, width=70, height=15)
response_text.pack(padx=10, pady=10)
response_text.insert(tk.END, "歡迎！輸入文字開始對話。\n\n")

# 使用者輸入框
user_input_text = tk.Text(root, height=3, width=70)
user_input_text.pack(padx=10, pady=5)

# 生成回應按鈕
generate_button = tk.Button(root, text="生成回應", command=generate_response)
generate_button.pack(pady=5)

# 資料清理按鈕
def process_data():
    if prepare_dataset():
        response_text.insert(tk.END, "資料處理完成，已生成 cleaned_data 資料夾內的 PDF、WORD 和 TXT 檔案！\n")

process_button = tk.Button(root, text="處理訓練資料", command=process_data)
process_button.pack(pady=5)

# 啟動應用
def start_app():
    response_text.insert(tk.END, "正在載入模型，請稍候...\n")
    root.update()
    try:
        load_model()
        response_text.insert(tk.END, "模型載入完成！開始對話吧！\n\n")
    except Exception as e:
        response_text.insert(tk.END, f"模型載入失敗：{str(e)}\n\n")

# 啟動應用
start_app()
root.mainloop()
